"""WordParser — extract structured content from .docx files.

Used by PptxGenerator and future document reformatting features.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)

# Mapping from python-docx style names to heading levels
_HEADING_STYLES: dict[str, int] = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "heading 1": 1,
    "heading 2": 2,
    "heading 3": 3,
    "Title": 1,
    "标题 1": 1,
    "标题 2": 2,
    "标题 3": 3,
}


@dataclass
class Section:
    """A section of a document with an optional heading and paragraphs."""

    level: int  # 0=no heading, 1-3=heading levels
    heading: str
    paragraphs: list[str] = field(default_factory=list)


@dataclass
class WordStructure:
    """Parsed structure of a Word document."""

    title: str
    sections: list[Section]
    total_paragraphs: int


class WordParseError(Exception):
    """Error raised when parsing or validating a Word document."""

    def __init__(self, message: str, reason: str) -> None:
        super().__init__(message)
        self.reason = reason  # not_found | unsupported | corrupted | empty


class WordParser:
    """Extract structured content from .docx files."""

    def validate(self, file_path: Path) -> None:
        """Validate a .docx file before parsing.

        Raises:
            WordParseError: If the file is invalid (wrong format, missing,
                corrupted, or empty).
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise WordParseError(f"File not found: {file_path}", reason="not_found")

        if file_path.suffix.lower() != ".docx":
            raise WordParseError(
                f"Unsupported file format: {file_path.suffix}",
                reason="unsupported",
            )

        try:
            doc = Document(str(file_path))
        except Exception as exc:
            raise WordParseError(f"Failed to open document: {exc}", reason="corrupted") from exc

        if not doc.paragraphs:
            raise WordParseError("Document contains no paragraphs", reason="empty")

    def parse(self, file_path: Path) -> WordStructure:
        """Parse a .docx file and return its structure.

        Args:
            file_path: Path to the .docx file.

        Returns:
            WordStructure with title, sections, and paragraph count.

        Raises:
            WordParseError: If the file is invalid.
        """
        file_path = Path(file_path)
        self.validate(file_path)

        doc = Document(str(file_path))

        sections: list[Section] = []
        title = ""
        total_paragraphs = 0

        for para in doc.paragraphs:
            text = para.text or ""
            style_name = para.style.name if para.style else ""

            heading_level = _HEADING_STYLES.get(style_name, 0)

            if heading_level > 0:
                # This paragraph is a heading
                if not title:
                    title = text
                sections.append(Section(level=heading_level, heading=text))
            else:
                # Non-heading paragraph
                total_paragraphs += 1
                if sections:
                    sections[-1].paragraphs.append(text)
                else:
                    # Paragraphs before any heading -> create a level-0 section
                    if not sections:
                        sections.append(Section(level=0, heading=""))
                    sections[-1].paragraphs.append(text)

        if not title:
            title = file_path.stem

        if not sections:
            # Edge case: document has paragraphs but all were headings
            # (handled by validate, but defensive)
            sections.append(Section(level=0, heading=""))

        return WordStructure(
            title=title,
            sections=sections,
            total_paragraphs=total_paragraphs,
        )
