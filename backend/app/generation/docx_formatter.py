from __future__ import annotations

import logging
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

logger = logging.getLogger(__name__)

FONT_CANDIDATES = {
    "title": ("方正小标宋简体", "宋体"),
    "heading1": ("方正黑体_GBK", "黑体"),
    "heading2": ("方正楷体_GBK", "楷体"),
    "heading3": ("方正仿宋_GBK", "仿宋"),
    "body": ("方正仿宋_GBK", "仿宋"),
    "signature": ("方正仿宋_GBK", "仿宋"),
    "recipient": ("方正仿宋_GBK", "仿宋"),
}

FONT_SIZES = {
    "title": Pt(22),
    "heading1": Pt(16),
    "heading2": Pt(16),
    "heading3": Pt(16),
    "body": Pt(16),
    "signature": Pt(16),
    "recipient": Pt(16),
}

ALIGNMENTS = {
    "title": WD_ALIGN_PARAGRAPH.CENTER,
    "heading1": WD_ALIGN_PARAGRAPH.LEFT,
    "heading2": WD_ALIGN_PARAGRAPH.LEFT,
    "heading3": WD_ALIGN_PARAGRAPH.LEFT,
    "body": WD_ALIGN_PARAGRAPH.LEFT,
    "signature": WD_ALIGN_PARAGRAPH.RIGHT,
    "recipient": WD_ALIGN_PARAGRAPH.LEFT,
}

EN_NUM_PATTERN = re.compile(r"[a-zA-Z0-9]+")

# Chinese government document heading patterns
CN_HEADING1 = re.compile(r"^[一二三四五六七八九十百]+、")
CN_HEADING2 = re.compile(r"^[（(][一二三四五六七八九十百]+[）)]")
NUM_HEADING = re.compile(r"^\d+[.、]\s*")
DATE_PATTERN = re.compile(r"[\dX]{2,4}年[\dX]{1,2}月[\dX]{1,2}日")
HAS_SENTENCE_END = re.compile(r"[。；？！]")
RECIPIENT_END = re.compile(r"[：:]$")


class DocxFormatter:
    def __init__(self, output_dir: Path) -> None:
        self._output_dir = output_dir
        self._output_dir.mkdir(parents=True, exist_ok=True)
        self._font_map = self._detect_fonts()

    def format(self, content: str, doc_type: str, topic: str) -> Path:
        structure = self._parse_structure(content)
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(3.7)
            section.bottom_margin = Cm(3.5)
            section.left_margin = Cm(2.8)
            section.right_margin = Cm(2.6)

        for item in structure:
            style_type = item["type"]
            text = item["text"]
            font_name = self._font_map.get(style_type, self._font_map["body"])
            font_size = FONT_SIZES.get(style_type, FONT_SIZES["body"])
            alignment = ALIGNMENTS.get(style_type, ALIGNMENTS["body"])

            p = doc.add_paragraph()
            p.alignment = alignment
            pf = p.paragraph_format
            pf.line_spacing = Pt(28)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            if style_type in ("title", "recipient"):
                pass
            elif style_type == "signature":
                pf.right_indent = Pt(64)
            else:
                pf.first_line_indent = Pt(32)

            if style_type == "heading2" and "body_part" in item:
                self._add_mixed_heading2(p, item, font_size)
            else:
                self._add_text_with_font(p, text, font_name, font_size)

        filename = self._make_filename(doc_type, topic)
        path = self._output_dir / filename
        doc.save(str(path))
        return path

    def _parse_structure(self, content: str) -> list[dict]:
        raw_lines: list[str] = []
        for line in content.strip().split("\n"):
            line = line.strip()
            if line:
                raw_lines.append(self._strip_markdown(line))

        if not raw_lines:
            return []

        structure: list[dict] = []

        # First non-empty line → title (unless it matches a heading pattern)
        first = raw_lines[0]
        if first.startswith("# ") and not first.startswith("## "):
            structure.append({"type": "title", "text": first[2:]})
        elif first.startswith("## "):
            structure.append({"type": "title", "text": first[3:]})
        elif not CN_HEADING1.match(first) and not CN_HEADING2.match(first) and not NUM_HEADING.match(first):
            structure.append({"type": "title", "text": first})
        else:
            structure.append({"type": "body", "text": first})

        for line in raw_lines[1:]:
            if line.startswith("# ") and not line.startswith("## "):
                structure.append({"type": "title", "text": line[2:]})
            elif line.startswith("## "):
                structure.append({"type": "heading1", "text": line[3:]})
            elif line.startswith("### "):
                structure.append({"type": "heading2", "text": line[4:]})
            elif line.startswith("#### "):
                structure.append({"type": "heading3", "text": line[5:]})
            elif CN_HEADING1.match(line):
                structure.append({"type": "heading1", "text": line})
            elif CN_HEADING2.match(line):
                idx = line.find("。")
                if idx >= 0 and idx < len(line) - 1:
                    structure.append({
                        "type": "heading2",
                        "text": line,
                        "heading_part": line[: idx + 1],
                        "body_part": line[idx + 1 :],
                    })
                else:
                    structure.append({"type": "heading2", "text": line})
            elif NUM_HEADING.match(line):
                structure.append({"type": "heading3", "text": line})
            else:
                structure.append({"type": "body", "text": line})

        self._mark_signatures(structure)
        self._mark_recipient(structure)
        return structure

    def _mark_signatures(self, structure: list[dict]) -> None:
        for i in range(len(structure) - 1, -1, -1):
            if structure[i]["type"] != "body":
                continue
            text = structure[i]["text"]
            if DATE_PATTERN.search(text):
                structure[i]["type"] = "signature"
                if i > 0 and structure[i - 1]["type"] == "body":
                    prev = structure[i - 1]["text"]
                    if len(prev) < 20 and not HAS_SENTENCE_END.search(prev):
                        structure[i - 1]["type"] = "signature"
                break

    def _mark_recipient(self, structure: list[dict]) -> None:
        for i, item in enumerate(structure):
            if item["type"] == "title" and i + 1 < len(structure):
                next_item = structure[i + 1]
                if next_item["type"] == "body" and RECIPIENT_END.search(next_item["text"]):
                    next_item["type"] = "recipient"
                break

    def _add_mixed_heading2(self, paragraph, item: dict, font_size) -> None:
        heading_font = self._font_map.get("heading2", self._font_map["body"])
        body_font = self._font_map["body"]
        self._add_text_with_font(paragraph, item["heading_part"], heading_font, font_size)
        self._add_text_with_font(paragraph, item["body_part"], body_font, font_size)

    def _strip_markdown(self, text: str) -> str:
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"__(.+?)__", r"\1", text)
        text = re.sub(r"_(.+?)_", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        text = re.sub(r"~~(.+?)~~", r"\1", text)
        return text

    def _add_text_with_font(self, paragraph, text: str, font_name: str, font_size) -> None:
        parts = EN_NUM_PATTERN.split(text)
        en_parts = EN_NUM_PATTERN.findall(text)
        for i, run_text in enumerate(parts):
            if run_text:
                run = paragraph.add_run(run_text)
                run.font.name = font_name
                run.font.size = font_size
                run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            if i < len(en_parts):
                run = paragraph.add_run(en_parts[i])
                run.font.name = "Times New Roman"
                run.font.size = font_size
                run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _make_filename(self, doc_type: str, topic: str) -> str:
        safe_doc_type = re.sub(r"[^\w]", "", doc_type)[:20]
        safe_topic = re.sub(r"[^\w\u4e00-\u9fff]", "", topic)[:20]
        return f"{safe_doc_type}_{safe_topic}_{date.today().isoformat()}.docx"

    def _detect_fonts(self) -> dict:
        available = set()
        from platform import system

        if system() == "Darwin":
            import subprocess

            try:
                result = subprocess.run(
                    ["system_profiler", "SPFontsDataType"],
                    capture_output=True,
                    text=True,
                    timeout=30,
                )
                for names in FONT_CANDIDATES.values():
                    for n in names:
                        if n in result.stdout:
                            available.add(n)
            except Exception:
                pass

        resolved = {}
        for key, (preferred, fallback) in FONT_CANDIDATES.items():
            if preferred in available:
                resolved[key] = preferred
            else:
                resolved[key] = fallback
                if preferred != fallback:
                    logger.info("字体 %s 不可用，降级为 %s", preferred, fallback)
        return resolved
