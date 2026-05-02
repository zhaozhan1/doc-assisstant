from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.generation.word_parser import Section, WordParseError, WordParser, WordStructure


def _create_docx(path: Path, paragraphs: list[tuple[str, str]]) -> Path:
    """Create a test .docx file.

    Args:
        path: Output file path.
        paragraphs: List of (style_name, text) tuples.

    Returns:
        The path to the created file.
    """
    doc = Document()
    for style, text in paragraphs:
        doc.add_paragraph(text, style=style)
    doc.save(str(path))
    return path


@pytest.fixture
def parser() -> WordParser:
    return WordParser()


class TestWordParserParse:
    """Tests for WordParser.parse()."""

    def test_normal_document_with_headings(self, parser: WordParser, tmp_path: Path) -> None:
        """Heading 1 + Heading 2 + paragraphs -> correct sections."""
        docx_path = _create_docx(
            tmp_path / "normal.docx",
            [
                ("Heading 1", "一、总体要求"),
                (None, "第一段内容。"),
                ("Heading 2", "（一）具体要求"),
                (None, "第二段内容。"),
                (None, "第三段内容。"),
                ("Heading 1", "二、工作安排"),
                (None, "第四段内容。"),
            ],
        )

        result = parser.parse(docx_path)

        assert isinstance(result, WordStructure)
        assert result.title == "一、总体要求"
        assert len(result.sections) == 3
        assert result.total_paragraphs == 4  # non-heading paragraphs only

        # Section 1: Heading 1
        assert result.sections[0].level == 1
        assert result.sections[0].heading == "一、总体要求"
        assert result.sections[0].paragraphs == ["第一段内容。"]

        # Section 2: Heading 2 (nested under section 1)
        assert result.sections[1].level == 2
        assert result.sections[1].heading == "（一）具体要求"
        assert result.sections[1].paragraphs == ["第二段内容。", "第三段内容。"]

        # Section 3: Heading 1
        assert result.sections[2].level == 1
        assert result.sections[2].heading == "二、工作安排"
        assert result.sections[2].paragraphs == ["第四段内容。"]

    def test_plain_text_no_headings(self, parser: WordParser, tmp_path: Path) -> None:
        """Document with no headings -> single Section with level=0."""
        docx_path = _create_docx(
            tmp_path / "plain.docx",
            [
                (None, "纯文本段落一。"),
                (None, "纯文本段落二。"),
                (None, "纯文本段落三。"),
            ],
        )

        result = parser.parse(docx_path)

        assert result.title == "plain"  # filename stem
        assert len(result.sections) == 1
        assert result.sections[0].level == 0
        assert result.sections[0].heading == ""
        assert result.sections[0].paragraphs == [
            "纯文本段落一。",
            "纯文本段落二。",
            "纯文本段落三。",
        ]
        assert result.total_paragraphs == 3

    def test_title_from_first_heading(self, parser: WordParser, tmp_path: Path) -> None:
        """Title is the text of the first heading."""
        docx_path = _create_docx(
            tmp_path / "titled.docx",
            [
                ("Heading 1", "重要通知"),
                (None, "内容。"),
            ],
        )

        result = parser.parse(docx_path)
        assert result.title == "重要通知"

    def test_title_fallback_to_filename(self, parser: WordParser, tmp_path: Path) -> None:
        """No headings -> title falls back to filename stem."""
        docx_path = _create_docx(
            tmp_path / "我的文档.docx",
            [(None, "内容。")],
        )

        result = parser.parse(docx_path)
        assert result.title == "我的文档"

    def test_heading_3_level(self, parser: WordParser, tmp_path: Path) -> None:
        """Heading 3 is correctly parsed as level 3."""
        docx_path = _create_docx(
            tmp_path / "h3.docx",
            [
                ("Heading 1", "主标题"),
                ("Heading 2", "副标题"),
                ("Heading 3", "三级标题"),
                (None, "内容。"),
            ],
        )

        result = parser.parse(docx_path)
        assert result.sections[2].level == 3

    def test_total_paragraphs_counts_correctly(self, parser: WordParser, tmp_path: Path) -> None:
        """total_paragraphs counts all non-heading paragraphs."""
        docx_path = _create_docx(
            tmp_path / "count.docx",
            [
                ("Heading 1", "标题"),
                (None, "段一"),
                (None, "段二"),
                ("Heading 2", "子标题"),
                (None, "段三"),
                (None, ""),
            ],
        )

        result = parser.parse(docx_path)
        # Empty paragraphs are still counted
        assert result.total_paragraphs == 4


class TestWordParserValidate:
    """Tests for WordParser.validate()."""

    def test_rejects_txt_file(self, parser: WordParser, tmp_path: Path) -> None:
        """validate() rejects .txt files -> WordParseError reason='unsupported'."""
        txt_path = tmp_path / "test.txt"
        txt_path.write_text("hello")

        with pytest.raises(WordParseError) as exc_info:
            parser.validate(txt_path)

        assert exc_info.value.reason == "unsupported"

    def test_rejects_nonexistent_file(self, parser: WordParser, tmp_path: Path) -> None:
        """validate() rejects missing files -> WordParseError reason='not_found'."""
        missing = tmp_path / "missing.docx"

        with pytest.raises(WordParseError) as exc_info:
            parser.validate(missing)

        assert exc_info.value.reason == "not_found"

    def test_rejects_empty_docx(self, parser: WordParser, tmp_path: Path) -> None:
        """validate() rejects empty .docx -> WordParseError reason='empty'."""
        docx_path = _create_docx(tmp_path / "empty.docx", [])

        with pytest.raises(WordParseError) as exc_info:
            parser.validate(docx_path)

        assert exc_info.value.reason == "empty"

    def test_rejects_corrupted_file(self, parser: WordParser, tmp_path: Path) -> None:
        """validate() rejects corrupted files -> WordParseError reason='corrupted'."""
        corrupted = tmp_path / "corrupt.docx"
        corrupted.write_bytes(b"not a real docx file content")

        with pytest.raises(WordParseError) as exc_info:
            parser.validate(corrupted)

        assert exc_info.value.reason == "corrupted"

    def test_accepts_valid_docx(self, parser: WordParser, tmp_path: Path) -> None:
        """validate() does not raise for valid .docx files."""
        docx_path = _create_docx(
            tmp_path / "valid.docx",
            [(None, "some content")],
        )

        # Should not raise
        parser.validate(docx_path)


class TestSection:
    """Tests for Section dataclass."""

    def test_section_fields(self) -> None:
        section = Section(level=1, heading="Test", paragraphs=["a", "b"])
        assert section.level == 1
        assert section.heading == "Test"
        assert section.paragraphs == ["a", "b"]


class TestWordParseError:
    """Tests for WordParseError exception."""

    def test_error_has_reason(self) -> None:
        err = WordParseError("something went wrong", reason="corrupted")
        assert str(err) == "something went wrong"
        assert err.reason == "corrupted"
